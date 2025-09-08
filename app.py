"""
Streamlit App: Document → English Translator (Groq)
-------------------------------------------------
Quick start:
1) Save this file as `app.py`
2) `pip install -r requirements.txt`  (see inline list below)
3) Set your Groq key via either:
   - Streamlit Secrets:  `.streamlit/secrets.toml` with GROQ_API_KEY="your_key"
   - or a `.env` file with GROQ_API_KEY=your_key
4) Run: `streamlit run app.py`

Requirements (add to requirements.txt):
- streamlit>=1.37
- groq>=0.11.0
- PyMuPDF==1.24.9        # for PDFs (text-based)
- python-docx>=1.1.0      # to export .docx
- python-dotenv>=1.0.1    # optional, for .env

Notes:
- This app extracts text from TXT, DOCX, and (select) PDFs. If a PDF is scanned (image-only), you’ll see an OCR warning. For OCR, consider adding Tesseract or EasyOCR offline; not enabled here to keep Streamlit Cloud-friendly.
- Translation uses Groq LLM (default: llama-3.1-70b-versatile). You can switch models in the sidebar.
"""

from __future__ import annotations
import os
import io
import time
import textwrap
from typing import List, Tuple

import streamlit as st
from dotenv import load_dotenv

# File handlers
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None

# Groq client
try:
    from groq import Groq
except Exception:
    Groq = None

# -----------------------------
# Setup
# -----------------------------
st.set_page_config(page_title="📄→🇬🇧 Document Translator (Groq)", layout="wide")
st.title("📄→🇬🇧 Document Translator (Groq)")

with st.expander("How it works"):
    st.markdown(
        """
        **Upload** a `.pdf`, `.docx`, or `.txt` file. The app extracts the text and sends it in safe chunks
        to a Groq LLM with a **translation-focused** prompt. The output is stitched together and can be
        downloaded as **TXT** or **DOCX**; headings/tables are preserved as Markdown when possible.
        """
    )

# Secrets/Env
load_dotenv()
API_KEY = st.secrets.get("GROQ_API_KEY", os.getenv("GROQ_API_KEY", ""))

# Sidebar controls
with st.sidebar:
    st.header("Settings")
    if not API_KEY:
        st.warning("Add your GROQ_API_KEY to Streamlit Secrets or .env.")

    model = st.selectbox(
        "Groq model",
        options=[
            "llama-3.1-70b-versatile",
            "llama-3.1-8b-instant",
            "mixtral-8x7b-32768",
        ],
        index=0,
        help="70B for best quality; 8B for speed; Mixtral for long contexts.",
    )

    max_chars_per_chunk = st.slider(
        "Max characters per chunk",
        min_value=1500,
        max_value=8000,
        value=4000,
        step=500,
        help="Text is chunked before translation to avoid context limits.",
    )

    temperature = st.slider("Creativity (temperature)", 0.0, 1.0, 0.2, 0.1)
    add_qa_pass = st.checkbox(
        "Optional: add QA pass to correct names, numbers, and dates",
        value=True,
        help="Runs an extra quick pass to catch obvious errors."
    )

# -----------------------------
# Utilities
# -----------------------------

def _read_txt(file) -> str:
    data = file.read()
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def _read_docx(file) -> str:
    if Document is None:
        raise RuntimeError("python-docx not installed.")
    fbytes = io.BytesIO(file.read())
    doc = Document(fbytes)
    paras = [p.text for p in doc.paragraphs]
    text = "\n".join([p for p in paras if p is not None])
    return text.strip()


def _read_pdf(file) -> Tuple[str, int]:
    """Return (text, page_count). Requires PyMuPDF."""
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) not installed.")
    fbytes = io.BytesIO(file.read())
    doc = fitz.open(stream=fbytes, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        try:
            pages.append(page.get_text("text"))
        except Exception:
            pages.append("")
    text = "\n\n".join(pages).strip()
    return text, len(doc)


def extract_text(uploaded_file) -> Tuple[str, str]:
    """Return (text, source_info) for supported file types."""
    if uploaded_file is None:
        return "", ""
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return _read_txt(uploaded_file), "TXT"
    elif name.endswith(".docx"):
        return _read_docx(uploaded_file), "DOCX"
    elif name.endswith(".pdf"):
        text, pages = _read_pdf(uploaded_file)
        return text, f"PDF ({pages} pages)"
    else:
        raise ValueError("Unsupported file type. Please upload .pdf, .docx, or .txt")


def chunk_text(text: str, max_chars: int) -> List[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks, buf = [], ""
    for p in paragraphs:
        # if a paragraph is too large, wrap it conservatively
        if len(p) > max_chars:
            wrapped = textwrap.wrap(p, width=max_chars, replace_whitespace=False, drop_whitespace=False)
            for w in wrapped:
                if len(buf) + len(w) + 2 > max_chars:
                    if buf:
                        chunks.append(buf)
                    buf = w
                else:
                    buf += ("\n\n" if buf else "") + w
        else:
            if len(buf) + len(p) + 2 > max_chars:
                if buf:
                    chunks.append(buf)
                buf = p
            else:
                buf += ("\n\n" if buf else "") + p
    if buf:
        chunks.append(buf)
    return chunks


def groq_client(api_key: str):
    if Groq is None:
        raise RuntimeError("groq package not installed.")
    return Groq(api_key=api_key)


SYSTEM_TRANSLATOR = (
    "You are a professional legal/financial translator. Translate the user's text into natural, clear,\n"
    "Oxford-style English. Preserve meaning precisely. Keep numbers, names, and dates accurate.\n"
    "Maintain structure: use Markdown for headings, lists, and simple tables when helpful.\n"
    "Do not add commentary or summaries—only the translation."
)

SYSTEM_QA = (
    "You are a meticulous proofreader. Check the translation for mistakes in numbers, dates, names,\n"
    "and units. Correct obvious translation errors while preserving meaning. Return the corrected text only."
)


def translate_chunk(cli, model: str, src_text: str, temperature: float = 0.2) -> str:
    resp = cli.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_TRANSLATOR},
            {"role": "user", "content": src_text},
        ],
        temperature=temperature,
    )
    return resp.choices[0].message.content.strip()


def qa_pass(cli, model: str, translated_text: str) -> str:
    resp = cli.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_QA},
            {"role": "user", "content": translated_text},
        ],
        temperature=0.0,
    )
    return resp.choices[0].message.content.strip()


# -----------------------------
# Main UI
# -----------------------------
col_left, col_right = st.columns([1, 1])
with col_left:
    uploaded = st.file_uploader(
        "Upload document (.pdf, .docx, .txt)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=False,
    )

with col_right:
    sample_hint = st.checkbox("Show sample input/expected output format", value=False)
    if sample_hint:
        st.code(
            """# Example: original (Arabic, German, Urdu, etc.)\n- Keep paragraphs.\n- Maintain numbering and bullets.\n\n# Example: output (English)\n- Clean, professional English.\n- Headings and lists preserved using Markdown.\n""",
            language="markdown",
        )

start = st.button("Translate to English", type="primary", disabled=(uploaded is None))

if start and uploaded is not None:
    if not API_KEY:
        st.error("Missing GROQ_API_KEY. Add it to Streamlit Secrets or .env and rerun.")
        st.stop()

    # Extract text
    with st.status("Extracting text…", expanded=False) as status:
        try:
            text, src_kind = extract_text(uploaded)
            status.update(label=f"Extracted from {src_kind}")
        except Exception as e:
            st.exception(e)
            st.stop()

    if not text or len(text.strip()) < 20:
        st.warning(
            "Could not extract meaningful text. If this is a scanned PDF, enable OCR externally and try again."
        )
        st.stop()

    # Show a small preview of the source
    with st.expander("Preview: Original text (first 1,000 chars)"):
        st.text(text[:1000])

    # Chunk
    chunks = chunk_text(text, max_chars_per_chunk)
    st.info(f"Processing {len(chunks)} chunk(s)…")

    # Translate
    cli = groq_client(API_KEY)

    progress = st.progress(0)
    out_parts: List[str] = []
    errors: List[Tuple[int, str]] = []

    for i, ch in enumerate(chunks, start=1):
        try:
            translated = translate_chunk(cli, model, ch, temperature)
            out_parts.append(translated)
        except Exception as e:
            errors.append((i, str(e)))
            out_parts.append(f"\n[Translation error in chunk {i}: {e}]\n")
        progress.progress(i / len(chunks))
        # be polite to rate limits
        time.sleep(0.05)

    full_translation = "\n\n".join(out_parts).strip()

    if add_qa_pass and full_translation:
        with st.status("QA pass (numbers, names, dates)…", expanded=False):
            try:
                full_translation = qa_pass(cli, model, full_translation)
            except Exception as e:
                st.warning(f"QA pass failed: {e}")

    # Output preview
    st.subheader("✅ Translation Complete")
    st.caption("Rendered below (Markdown). Download as TXT or DOCX using the buttons.")
    st.markdown(full_translation)

    # Downloads
    txt_bytes = full_translation.encode("utf-8")
    st.download_button("⬇️ Download .txt", data=txt_bytes, file_name="translation_en.txt", mime="text/plain")

    if Document is not None:
        # Write a basic docx with paragraphs split by blank lines
        doc = Document()
        for block in full_translation.split("\n\n"):
            for line in block.split("\n"):
                doc.add_paragraph(line)
            doc.add_paragraph("")
        bio = io.BytesIO()
        doc.save(bio)
        st.download_button(
            "⬇️ Download .docx",
            data=bio.getvalue(),
            file_name="translation_en.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # Any issues?
    if errors:
        with st.expander("Errors encountered (safe to ignore if output looks fine)"):
            for idx, err in errors:
                st.write(f"Chunk {idx}: {err}")

# -----------------------------
# Footer
# -----------------------------
st.markdown(
    """
    ---
    **Tips**
    - For large PDFs, try increasing the chunk size cautiously.
    - If your PDF is scanned/image-based, pre-run OCR (Tesseract, cloud OCR) to extract text first.
    - To localize to *other* target languages, adjust the system translator message.
    - For enterprise controls, consider saving outputs to S3/Drive and adding audit metadata.
    """
)
